# Copyright GraphCaster. All Rights Reserved.

from __future__ import annotations

from pathlib import Path
from typing import AsyncIterator

from graph_caster.rag.loaders.base import Document, DocumentLoader


class DocxLoader(DocumentLoader):
    """Load a .docx file — one Document per heading-delimited section.

    Requires the optional ``rag-loaders-docx`` extra::

        pip install "graph-caster[rag-loaders-docx]"
    """

    def __init__(self, path: str | Path) -> None:
        self._path = Path(path)

    def _get_document(self):
        try:
            from docx import Document as DocxDocument  # type: ignore[import-untyped]
        except ImportError as exc:
            raise ImportError(
                "python-docx is required for DocxLoader. "
                'Install it with: pip install "graph-caster[rag-loaders-docx]"'
            ) from exc
        return DocxDocument(str(self._path))

    async def load(self) -> list[Document]:
        return [doc async for doc in self.lazy_load()]

    async def lazy_load(self) -> AsyncIterator[Document]:
        docx_doc = self._get_document()
        source = str(self._path)

        current_heading: str | None = None
        current_lines: list[str] = []

        def _flush(heading: str | None, lines: list[str]):
            text = "\n".join(lines).strip()
            if text:
                return Document(
                    page_content=text,
                    metadata={"source": source, "section": heading or ""},
                )
            return None

        for para in docx_doc.paragraphs:
            style_name = para.style.name if para.style else ""
            is_heading = style_name.lower().startswith("heading")
            if is_heading:
                flushed = _flush(current_heading, current_lines)
                if flushed is not None:
                    yield flushed
                current_heading = para.text.strip()
                current_lines = []
            else:
                if para.text.strip():
                    current_lines.append(para.text)

        flushed = _flush(current_heading, current_lines)
        if flushed is not None:
            yield flushed
        elif not current_lines and current_heading is not None:
            yield Document(
                page_content=current_heading,
                metadata={"source": source, "section": current_heading},
            )
